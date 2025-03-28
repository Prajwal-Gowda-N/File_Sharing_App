from pdf2docx import Converter
from docx import Document
import io
import os
from reportlab.pdfgen import canvas
from io import BytesIO

# Function to convert Word to PDF
def convert_word_to_pdf(word_file):
    # Save the incoming file to a temporary location
    temp_file_path = "temp_word_file.docx"
    with open(temp_file_path, 'wb') as f:
        f.write(word_file.read())

    output_pdf = "output.pdf"
    # Use the library or any logic to convert DOCX to PDF here.
    # This is a simplified approach using the ReportLab library, which can be improved as needed
    c = canvas.Canvas(output_pdf)
    c.drawString(100, 750, "Word to PDF Conversion!")
    c.save()

    with open(output_pdf, 'rb') as f:
        pdf_data = f.read()
    
    os.remove(temp_file_path)  # Cleanup the temporary file
    os.remove(output_pdf)      # Remove generated PDF file

    return pdf_data


# Function to convert PDF to Word
def convert_pdf_to_word(pdf_file):
    # Temporary saving the uploaded PDF
    temp_pdf_path = "temp_pdf_file.pdf"
    with open(temp_pdf_path, 'wb') as f:
        f.write(pdf_file.read())

    # Convert the PDF to Word (simplified here, use actual conversion logic)
    doc = Document()
    doc.add_paragraph("Converted PDF to Word Content")
    doc_file_path = "converted.docx"
    doc.save(doc_file_path)

    with open(doc_file_path, 'rb') as f:
        word_data = f.read()

    os.remove(temp_pdf_path)  # Cleanup the temporary file
    os.remove(doc_file_path)  # Remove the generated Word file

    return word_data


import torch
from transformers import BartForConditionalGeneration, BartTokenizer
import pdfplumber
import docx
from io import BytesIO

# Load the BART model and tokenizer from Hugging Face
model_name = "facebook/bart-large-cnn"
tokenizer = BartTokenizer.from_pretrained(model_name)
model = BartForConditionalGeneration.from_pretrained(model_name)

# Function to summarize text using the Hugging Face model
def summarize_text_huggingface(text, num_sentences=3):
    inputs = tokenizer([text], max_length=1024, return_tensors="pt", truncation=True, padding=True)
    summary_ids = model.generate(inputs["input_ids"], num_beams=4, min_length=20, max_length=150, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(BytesIO(pdf_file.read())) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_file):
    doc = docx.Document(BytesIO(docx_file.read()))
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

# Function to extract text from a TXT file
def extract_text_from_txt(txt_file):
    return txt_file.read().decode("utf-8")
